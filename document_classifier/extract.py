"""
Извлечение текста из PDF, Excel, Word, txt/csv и изображений для обучения/классификации.

Параметр ``text_extract_mode`` (``auto`` | ``local`` | ``ocr``):

- **PDF** — ``auto``/``ocr`` → Yandex Vision по растеризованным страницам; ``local`` — текстовый слой.
- **PNG/JPEG** — Yandex Vision (``local`` недоступен).
- **docx / xlsx / xls / txt / csv** — всегда **локально** (python-docx / pandas / чтение файла);
  режим извлечения на эти форматы не влияет.

Yandex Cloud: ``OCR_YANDEX_*`` (см. ``_ocr_bytes_yandex`` / ``_ocr_pdf_yandex``).

После классификации и извлечения реквизитов наименование поставщика для API
нормализуется в ``invoice_fields.supplier_display_name`` (срез типичного мусора OCR:
длинные числа, битые фрагменты региона, склейки строк).
"""
from __future__ import annotations

import base64
import os
import time
from pathlib import Path
from typing import Any, Optional

# Максимум символов на один документ (дальше обрезается при токенизации)
MAX_DOC_CHARS = 50_000

def resolve_readable_document(path: Path) -> Path:
    """
    Если передана папка — берётся первый подходящий файл (приоритет: pdf, xlsx, xls, затем png/jpg/jpeg).
    Иначе должен быть файл с поддерживаемым расширением.
    """
    path = Path(path).expanduser()
    if not path.exists():
        hint = ""
        par = path.parent
        if par.exists() and path.name:
            try:
                if path.suffix:
                    same = [p.name for p in par.iterdir() if p.is_file() and p.suffix.lower() == path.suffix.lower()]
                    if same:
                        hint = f" В этой папке есть файлы: {', '.join(same[:15])}"
                else:
                    sub = [p.name for p in par.iterdir()]
                    if sub:
                        hint = f" Содержимое «{par}»: {', '.join(sub[:15])}"
            except OSError:
                pass
        raise FileNotFoundError(
            f"Не найдено: {path}.{hint} "
            f"Проверьте имя папки (например, «СчетНаОплату») и что файл существует.",
        )

    if path.is_dir():
        found: list[Path] = []
        for ext in (".pdf", ".docx", ".xlsx", ".xls", ".png", ".jpg", ".jpeg"):
            found.extend(sorted(path.glob(f"*{ext}"), key=lambda p: p.name.lower()))
        if not found:
            raise ValueError(
                f"В папке нет поддерживаемых файлов (pdf, docx, xlsx, xls, png, jpg, jpeg): {path}",
            )
        return found[0]

    if not path.is_file():
        raise ValueError(f"Укажите файл или папку с документами: {path}")

    return path


def _extract_pdf_pymupdf(path: Path) -> str:
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            t = page.get_text("text") or ""
            if t.strip():
                parts.append(t.strip())
    return "\n\n".join(parts).strip()


def _extract_pdf_pypdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t.strip())
    return "\n\n".join(parts).strip()


def _normalize_text_extract_mode(mode: str) -> str:
    m = (mode or "auto").strip().lower()
    if m not in ("auto", "local", "ocr"):
        raise ValueError(
            f"text_extract_mode: ожидалось auto | local | ocr, получено {mode!r}",
        )
    return m


def _read_pdf_text_layer(path: Path) -> str:
    """Только текстовый слой PDF (PyMuPDF или pypdf), без OCR."""
    path = Path(path)
    try:
        text = _extract_pdf_pymupdf(path)
    except ImportError:
        try:
            text = _extract_pdf_pypdf(path)
        except ImportError as e:
            raise ImportError(
                "Нет библиотеки для PDF. Установите одну: pip install pypdf  или  pip install pymupdf",
            ) from e
    return (text or "").strip()


def _yandex_ocr_enabled() -> bool:
    return os.environ.get("OCR_YANDEX_ENABLED", "").strip().lower() in ("1", "true", "yes", "on")


def _yandex_vision_available() -> bool:
    """Папка + токен для recognizeText (как у JPEG)."""
    if not _yandex_ocr_enabled():
        return False
    if not (os.environ.get("OCR_YANDEX_FOLDER_ID") or "").strip():
        return False
    if (os.environ.get("OCR_YANDEX_IAM_TOKEN") or "").strip():
        return True
    if (os.environ.get("OCR_YANDEX_OAUTH_TOKEN") or "").strip():
        return True
    return False


def extract_text_from_pdf(path: Path, *, text_extract_mode: str = "auto") -> str:
    """
    Извлечение текста из PDF.

    - **auto** (по умолчанию): при настроенном Yandex Cloud (**OCR_YANDEX_***) **всегда**
      Vision OCR по страницам (тот же путь, что и для JPEG). Если OCR недоступен или пустой
      ответ — fallback на текстовый слой PDF.
    - **local**: только текстовый слой (PyMuPDF/pypdf), без облака.
    - **ocr**: только Yandex Vision; ошибка конфигурации/сети пробрасывается наверх.
    """
    path = Path(path)
    mode = _normalize_text_extract_mode(text_extract_mode)

    if mode == "local":
        text = _read_pdf_text_layer(path)
        return text[:MAX_DOC_CHARS] if text else ""

    if mode == "ocr":
        ocr = _ocr_pdf_yandex(path)
        return (ocr or "").strip()[:MAX_DOC_CHARS]

    # auto — предпочитаем Vision, как для изображений
    if _yandex_vision_available():
        try:
            ocr = _ocr_pdf_yandex(path)
            if ocr and ocr.strip():
                return ocr.strip()[:MAX_DOC_CHARS]
        except Exception:
            pass
    text = _read_pdf_text_layer(path)
    return text[:MAX_DOC_CHARS] if text else ""


_YANDEX_IAM_CACHE: dict[str, Any] = {"token": None, "expires_at": None}


def _yandex_iam_token() -> str:
    """
    Возвращает IAM-токен для Yandex Cloud.

    Варианты:
    - OCR_YANDEX_IAM_TOKEN: готовый IAM-токен (самый простой).
    - OCR_YANDEX_OAUTH_TOKEN: OAuth-токен Яндекса, будет обменян на IAM-токен.
    """
    direct = (os.environ.get("OCR_YANDEX_IAM_TOKEN") or "").strip()
    if direct:
        return direct

    oauth = (os.environ.get("OCR_YANDEX_OAUTH_TOKEN") or "").strip()
    if not oauth:
        raise RuntimeError(
            "Для OCR (Yandex) задайте OCR_YANDEX_IAM_TOKEN или OCR_YANDEX_OAUTH_TOKEN.",
        )

    # Кэшируем IAM-токен (обычно живёт несколько часов)
    now = time.time()
    cached = _YANDEX_IAM_CACHE.get("token")
    exp = _YANDEX_IAM_CACHE.get("expires_at")
    if cached and isinstance(exp, (int, float)) and now < exp - 60:
        return str(cached)

    import requests

    r = requests.post(
        "https://iam.api.cloud.yandex.net/iam/v1/tokens",
        json={"yandexPassportOauthToken": oauth},
        timeout=20,
    )
    r.raise_for_status()
    data = r.json()
    token = (data.get("iamToken") or "").strip()
    if not token:
        raise RuntimeError("Yandex IAM: пустой iamToken в ответе")

    expires_at = data.get("expiresAt")
    # expiresAt обычно ISO8601, но формат может меняться — безопасно ставим дефолт на 10 минут
    ttl = 600
    if isinstance(expires_at, str) and expires_at:
        try:
            # 2026-04-02T11:07:59Z
            from datetime import datetime, timezone

            dt = datetime.fromisoformat(expires_at.replace("Z", "+00:00"))
            ttl = max(60, int(dt.replace(tzinfo=timezone.utc).timestamp() - now))
        except Exception:
            ttl = 600

    _YANDEX_IAM_CACHE["token"] = token
    _YANDEX_IAM_CACHE["expires_at"] = now + ttl
    return token


def _yandex_extract_text_from_ocr_response(data: Any) -> str:
    """
    Достаёт распознанный текст из ответа Yandex Vision OCR.
    Формат ответа может отличаться; вытаскиваем максимально устойчиво.
    """
    if not isinstance(data, dict):
        return ""
    res = data.get("result")
    if isinstance(res, dict):
        ta = res.get("textAnnotation")
        if isinstance(ta, dict):
            for key in ("fullText", "text"):
                v = ta.get(key)
                if isinstance(v, str) and v.strip():
                    return v.strip()
        # Иногда текст лежит в blocks/lines/words
        blocks = res.get("blocks")
        if isinstance(blocks, list):
            parts: list[str] = []
            for b in blocks:
                if not isinstance(b, dict):
                    continue
                lines = b.get("lines")
                if not isinstance(lines, list):
                    continue
                for ln in lines:
                    if not isinstance(ln, dict):
                        continue
                    words = ln.get("words")
                    if isinstance(words, list) and words:
                        wtxt = []
                        for w in words:
                            if isinstance(w, dict):
                                t = w.get("text")
                                if isinstance(t, str) and t.strip():
                                    wtxt.append(t.strip())
                        if wtxt:
                            parts.append(" ".join(wtxt))
            if parts:
                return "\n".join(parts).strip()
    return ""


def _ocr_bytes_yandex(*, mime_type: str, content_bytes: bytes, language_codes: list[str]) -> str:
    """
    OCR изображения через Yandex Cloud Vision OCR.
    mime_type: PNG|JPEG
    """
    folder_id = (os.environ.get("OCR_YANDEX_FOLDER_ID") or "").strip()
    if not folder_id:
        raise RuntimeError("Для OCR (Yandex) задайте OCR_YANDEX_FOLDER_ID.")

    import requests

    token = _yandex_iam_token()
    url = "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {token}",
        "x-folder-id": folder_id,
    }
    payload = {
        "mimeType": mime_type,
        "languageCodes": language_codes,
        "model": "page",
        "content": base64.b64encode(content_bytes).decode("ascii"),
    }
    r = requests.post(url, headers=headers, json=payload, timeout=60)
    r.raise_for_status()
    return _yandex_extract_text_from_ocr_response(r.json())


def _yandex_language_codes() -> list[str]:
    langs = (os.environ.get("OCR_YANDEX_LANGS") or "ru,en").strip()
    codes = [x.strip() for x in langs.split(",") if x.strip()]
    return codes or ["ru", "en"]


def _ocr_pdf_yandex(path: Path) -> str:
    """
    OCR PDF через Yandex Cloud Vision OCR.

    Требует:
    - OCR_YANDEX_FOLDER_ID
    - OCR_YANDEX_IAM_TOKEN или OCR_YANDEX_OAUTH_TOKEN
    Опционально:
    - OCR_YANDEX_LANGS (по умолчанию: ru,en)
    - OCR_YANDEX_DPI (по умолчанию: 200)
    - OCR_YANDEX_MAX_PAGES (по умолчанию: 10)
    """
    folder_id = (os.environ.get("OCR_YANDEX_FOLDER_ID") or "").strip()
    if not folder_id:
        raise RuntimeError("Для OCR (Yandex) задайте OCR_YANDEX_FOLDER_ID.")

    language_codes = _yandex_language_codes()

    dpi = int(os.environ.get("OCR_YANDEX_DPI", "200"))
    dpi = min(400, max(72, dpi))
    max_pages = int(os.environ.get("OCR_YANDEX_MAX_PAGES", "10"))
    max_pages = min(50, max(1, max_pages))

    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ImportError("Для OCR PDF нужен pymupdf: pip install pymupdf") from e

    parts: list[str] = []
    with fitz.open(path) as doc:
        n = min(len(doc), max_pages)
        for i in range(n):
            page = doc.load_page(i)
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            img_bytes = pix.tobytes("png")
            text = _ocr_bytes_yandex(mime_type="PNG", content_bytes=img_bytes, language_codes=language_codes)
            if text:
                parts.append(text)

    return "\n\n".join(parts).strip()


def extract_text_from_image(path: Path) -> str:
    """
    Текст со скриншота/скана: только Yandex Cloud Vision OCR (PNG/JPEG).

    Требуются те же переменные, что для OCR PDF: OCR_YANDEX_FOLDER_ID,
    OCR_YANDEX_IAM_TOKEN или OCR_YANDEX_OAUTH_TOKEN.
    """
    path = Path(path)
    suf = path.suffix.lower()
    content = path.read_bytes()
    language_codes = _yandex_language_codes()
    if suf == ".png":
        mime = "PNG"
    elif suf in (".jpg", ".jpeg"):
        mime = "JPEG"
    else:
        raise ValueError(f"Изображение для OCR: ожидались .png, .jpg, .jpeg, получено {suf!r}")
    text = _ocr_bytes_yandex(mime_type=mime, content_bytes=content, language_codes=language_codes)
    return (text or "").strip()[:MAX_DOC_CHARS]


def extract_text_from_excel(path: Path) -> str:
    path = Path(path)
    suf = path.suffix.lower()
    try:
        import pandas as pd
    except ImportError as e:
        raise ImportError("Установите pandas openpyxl: pip install pandas openpyxl") from e

    rows: list[str] = []
    if suf == ".xlsx":
        xl = pd.ExcelFile(path, engine="openpyxl")
        for sheet in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet, header=None, dtype=str)
            rows.append(f"[Лист: {sheet}]")
            for _, r in df.iterrows():
                cells = [str(c) for c in r.tolist() if pd.notna(c) and str(c).strip()]
                if cells:
                    rows.append("\t".join(cells))
    elif suf == ".xls":
        try:
            xl = pd.ExcelFile(path, engine="xlrd")
        except ImportError as e:
            raise ImportError("Для .xls установите xlrd: pip install xlrd") from e
        for sheet in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet, header=None, dtype=str)
            rows.append(f"[Лист: {sheet}]")
            for _, r in df.iterrows():
                cells = [str(c) for c in r.tolist() if pd.notna(c) and str(c).strip()]
                if cells:
                    rows.append("\t".join(cells))
    else:
        raise ValueError(f"Неподдерживаемый Excel: {suf}")

    text = "\n".join(rows).strip()
    return text[:MAX_DOC_CHARS]


def extract_text_from_docx(path: Path) -> str:
    """Текст из Word (.docx): абзацы и ячейки таблиц."""
    path = Path(path)
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("Для .docx установите: pip install python-docx") from e

    document = Document(str(path))
    parts: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = "\t".join(cells)
            if line.strip():
                parts.append(line)
    text = "\n".join(parts).strip()
    return text[:MAX_DOC_CHARS]


def extract_text_from_file(path: Path, *, text_extract_mode: str = "auto") -> str:
    """
    ``text_extract_mode``: ``auto`` | ``local`` | ``ocr`` (для PDF и изображений).

    **docx / xlsx / xls / txt / csv** — всегда локальное чтение; параметр режима не используется.

    **Изображения**: ``local`` недоступен — только ``auto``/``ocr`` (Yandex OCR).

    **PDF**: ``auto``/``ocr`` — Yandex по страницам (растр); ``local`` — текстовый слой.
    """
    path = resolve_readable_document(Path(path))
    suf = path.suffix.lower()
    mode = _normalize_text_extract_mode(text_extract_mode)

    if suf == ".pdf":
        return extract_text_from_pdf(path, text_extract_mode=mode)

    if suf == ".docx":
        return extract_text_from_docx(path)

    if suf in (".xlsx", ".xls"):
        return extract_text_from_excel(path)

    if suf in (".txt", ".csv"):
        return path.read_text(encoding="utf-8", errors="replace").strip()[:MAX_DOC_CHARS]

    if suf in (".png", ".jpg", ".jpeg"):
        if mode == "local":
            raise ValueError(
                "Режим «только локально»: для PNG/JPEG текст без OCR недоступен. "
                "Выберите «Авто» или «Всегда OCR», либо загрузите PDF с текстовым слоем.",
            )
        return extract_text_from_image(path)
    ext = suf if suf else "(нет расширения)"
    raise ValueError(
        f"Неподдерживаемый формат «{ext}» для «{path.name}». "
        f"Нужны: .pdf, .docx, .xlsx, .xls, .txt, .csv, .png/.jpg/.jpeg (OCR Yandex) — или укажите папку с файлом.",
    )


def ocr_yandex_bytes_for_test(filename: str, content: bytes) -> str:
    """
    Утилита для тестирования токена/папки OCR Yandex из API.
    Только PDF (страницы → PNG) и изображения PNG/JPEG.
    """
    import tempfile

    name = filename or "upload"
    suf = Path(name).suffix.lower()
    language_codes = _yandex_language_codes()

    if suf == ".pdf":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(content)
            p = Path(tmp.name)
        try:
            return _ocr_pdf_yandex(p)
        finally:
            try:
                p.unlink(missing_ok=True)
            except OSError:
                pass

    if suf in (".png",):
        return _ocr_bytes_yandex(mime_type="PNG", content_bytes=content, language_codes=language_codes) or ""
    if suf in (".jpg", ".jpeg"):
        return _ocr_bytes_yandex(mime_type="JPEG", content_bytes=content, language_codes=language_codes) or ""

    raise ValueError("Для OCR теста поддерживаются только: .pdf, .png, .jpg, .jpeg")
